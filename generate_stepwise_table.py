"""
generate_stepwise_table.py  —  produce the two-panel stepwise results table.

Panel A: OC stepwise results (Table 1 in paper)
Panel B: OC vs ABM comparison at matched B values, showing Kirkwood indirect effect

Run any time the model or ABM results change.
Output: stepwise_table.docx  and  stepwise_table_data.json
"""
import json
import numpy as np
from model import Params, run_cell

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not available — JSON only")

# ── Panel A conditions ──────────────────────────────────────────────────────
CONDITIONS_A = [
    ("Baseline: B=0, no costs, flat g, no Kirkwood",
     dict(B=0, gm_a_zero=500, mm_scale=0, use_prod_curve=False,
          delta_rep=0, B_pgm=0)),
    ("+ grandmothering B=1, constant efficacy, no costs",
     dict(B=1, gm_a_zero=500, mm_scale=0, use_prod_curve=False,
          delta_rep=0, B_pgm=0)),
    ("+ declining grandmother efficacy",
     dict(B=1, gm_a_zero=90, mm_scale=0, use_prod_curve=False,
          delta_rep=0, B_pgm=0)),
    ("+ late-reproduction costs (q(a) + p_mat(a))",
     dict(B=1, gm_a_zero=90, mm_scale=1, use_prod_curve=False,
          delta_rep=0, B_pgm=0)),
    ("+ age-graded productivity g(a)  [Kaplan et al. 2000]",
     dict(B=1, gm_a_zero=90, mm_scale=1, use_prod_curve=True,
          delta_rep=0, B_pgm=0)),
    ("+ Kirkwood effort-dependent damage δ(u)  [DST]",
     dict(B=1, gm_a_zero=90, mm_scale=1, use_prod_curve=True,
          delta_rep=0.006, B_pgm=0)),
    ("+ second grandmother (B_pgm=1, p_u=0.05)",
     dict(B=1, gm_a_zero=90, mm_scale=1, use_prod_curve=True,
          delta_rep=0.006, B_pgm=1, p_u=0.05)),
]

# ── Panel B: OC vs ABM at matched B values (full model) ────────────────────
FULL_MODEL = dict(gm_a_zero=90, mm_scale=1, use_prod_curve=True,
                  delta_rep=0.006, B_pgm=0)

CONDITIONS_B_OC = [
    ("B=0  (no grandmothering)",             dict(B=0,  **FULL_MODEL)),
    ("B=1  MGM only",                         dict(B=1,  **FULL_MODEL)),
    ("B=1  MGM + PGM (p_u=0.05)",            dict(B=1,  **{**FULL_MODEL, "B_pgm":1, "p_u":0.05})),
    ("B=3  MGM only",                         dict(B=3,  **FULL_MODEL)),
    ("B=6  MGM only",                         dict(B=6,  **FULL_MODEL)),
]

# ABM long-run results from iMac (K=10,000, 100,000 yr, 5 seeds)
ABM_RESULTS = {
    "B=0  (no grandmothering)":        dict(Tr=47.7, Tr_sd=0.02, Td=48.7, PRLS=1.0, l70=0.20),
    "B=1  MGM only":                   dict(Tr=47.5, Tr_sd=0.03, Td=48.0, PRLS=0.5, l70=0.17),
    "B=1  MGM + PGM (p_u=0.05)":       dict(Tr=47.5, Tr_sd=0.03, Td=48.0, PRLS=0.5, l70=0.17),
    "B=3  MGM only":                   dict(Tr=47.1, Tr_sd=0.03, Td=48.3, PRLS=1.2, l70=0.17),
    "B=6  MGM only":                   dict(Tr=46.7, Tr_sd=0.07, Td=49.4, PRLS=2.7, l70=0.19),
}

GURVEN_KAPLAN_L70 = (0.40, 0.47)

Tstops = np.arange(30, 51)

def compute_oc_row(kw):
    W = np.array([run_cell(Params(a_ceiling=int(T), rate=0.0, **kw),
                           free_u=False)[0]["fitness"] for T in Tstops])
    W_ceil = W[-1]
    gain = float(100 * (W.max() - W_ceil) / abs(W_ceil))
    T_star = int(Tstops[np.argmax(W)])
    p = Params(rate=0.0, **kw)
    _, sim = run_cell(p, free_u=False)
    ages = np.arange(p.a_mat, p.a_max + 1)
    l70 = float(sim["l"][np.searchsorted(ages, 70)])
    return round(gain, 2), T_star, round(float(sim["PRLS"]), 1), round(l70, 2)

# ── Compute ─────────────────────────────────────────────────────────────────
print("Computing Panel A…")
rows_a = []
for label, kw in CONDITIONS_A:
    g, Ts, prls, l70 = compute_oc_row(kw)
    rows_a.append(dict(condition=label, Q1_gain=g, Q1_Tstar=Ts,
                       Q2_PRLS=prls, Q2_l70=l70))
    print(f"  {label[:54]:54s} {g:.2f}%  T*={Ts}  PRLS={prls}  l70={l70}")

print("\nComputing Panel B (OC side)…")
rows_b = []
for label, kw in CONDITIONS_B_OC:
    g, Ts, prls, l70 = compute_oc_row(kw)
    abm = ABM_RESULTS.get(label, {})
    rows_b.append(dict(
        condition=label,
        oc_Tstar=Ts, oc_gain=g, oc_PRLS=prls, oc_l70=l70,
        abm_Tr=abm.get("Tr"), abm_TrSD=abm.get("Tr_sd"),
        abm_Td=abm.get("Td"), abm_PRLS=abm.get("PRLS"), abm_l70=abm.get("l70")
    ))
    a = abm
    print(f"  {label}: OC T*={Ts} PRLS={prls} l70={l70} | "
          f"ABM Tr={a.get('Tr')}±{a.get('Tr_sd')} Td={a.get('Td')} "
          f"PRLS={a.get('PRLS')} l70={a.get('l70')}")

data = {"panel_A": rows_a, "panel_B": rows_b}
with open("stepwise_table_data.json", "w") as f:
    json.dump(data, f, indent=2)
print("\nSaved stepwise_table_data.json")

# ── Build Word document ──────────────────────────────────────────────────────
if not HAS_DOCX:
    print("No docx — done"); exit()

def shade_cell(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear"); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=9, align="center", fill=None):
    if fill: shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if align == "left"
                   else WD_ALIGN_PARAGRAPH.CENTER)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

HDR_FILL = "1b3a5c"; HDR_TXT = (255,255,255)
ALT1 = "EFF4FA"; ALT2 = "FFFFFF"
BLUE = (27,106,168); RED = (180,30,30); GREEN = (20,120,20)
GK_FILL = "e8f5e8"  # light green for rows meeting Gurven-Kaplan target

doc = Document()
for section in doc.sections:
    section.page_width = Cm(29.7); section.page_height = Cm(21.0)
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = section.bottom_margin = Cm(1.5)

# ═══════════════════════════════════════════════════════════════════════════
# PANEL A
# ═══════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Table 1A.  Stepwise OC predictions: how answers change as biology is added")
r.bold = True; r.font.size = Pt(11)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.add_run("B = 1 (maternal grandmother only), R0 stationary currency. "
          "Q1 = cessation-age selection;  Q2 = post-reproductive lifespan.  "
          "✓ = meets Gurven–Kaplan target l₇₀ ≥ 0.40.").font.size = Pt(9)
doc.add_paragraph()

HDRS_A = ["Model condition", "Q1\ngain (%)", "Q1\nT* (yr)", "Q2\nPRLS (yr)", "Q2\nl₇₀"]
CWID_A  = [11.5, 2.0, 2.5, 2.2, 1.8]
tbl = doc.add_table(rows=1+len(rows_a), cols=len(HDRS_A))
tbl.style = "Table Grid"
for i,(h,w) in enumerate(zip(HDRS_A, CWID_A)):
    c = tbl.rows[0].cells[i]; c.width = Cm(w)
    set_cell(c, h, bold=True, color=HDR_TXT, fill=HDR_FILL,
             align="left" if i==0 else "center")
for ri, row in enumerate(rows_a):
    tr = tbl.rows[ri+1]
    gk = row["Q2_l70"] >= GURVEN_KAPLAN_L70[0]
    fill = GK_FILL if gk else (ALT1 if ri%2==0 else ALT2)
    vals = [row["condition"],
            f'{row["Q1_gain"]:.2f}',
            str(row["Q1_Tstar"]),
            f'{row["Q2_PRLS"]:.1f}',
            f'{row["Q2_l70"]:.2f}' + (" ✓" if gk else "")]
    last = ri == len(rows_a)-1
    for i,(v,w) in enumerate(zip(vals, CWID_A)):
        c = tr.cells[i]; c.width = Cm(w)
        col = BLUE if (last and i>0) else None
        set_cell(c, v, bold=last and i>0, color=col, fill=fill,
                 align="left" if i==0 else "center")

doc.add_paragraph()
fn = doc.add_paragraph()
fn.add_run("Green shading: l₇₀ ≥ 0.40 (Gurven–Kaplan lower bound).  "
           "Bold blue: full model values.").font.size = Pt(8)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# PANEL B
# ═══════════════════════════════════════════════════════════════════════════
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Table 1B.  OC predictions vs ABM equilibria at matched B values (full model, all biology)")
r2.bold = True; r2.font.size = Pt(11)

s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.add_run("OC = drift-free optimal-control predictions.  "
           "ABM = long-run evolutionary equilibria (K=10,000, 100,000 yr, 5 seeds).  "
           "↓ marks ABM values below the B=0 ABM baseline (Kirkwood indirect effect).").font.size = Pt(9)
doc.add_paragraph()

HDRS_B = ["Condition",
           "OC: Q1\ngain (%)", "OC: T*\n(yr)", "OC: PRLS\n(yr)", "OC: l₇₀",
           "ABM: Tr\n(yr ± SD)", "ABM: Td\n(yr)", "ABM: PRLS\n(yr)", "ABM: l₇₀"]
CWID_B  = [5.2, 1.8, 1.6, 1.8, 1.5,   2.4, 1.6, 1.7, 1.5]

tbl2 = doc.add_table(rows=1+len(rows_b), cols=len(HDRS_B))
tbl2.style = "Table Grid"
for i,(h,w) in enumerate(zip(HDRS_B,CWID_B)):
    c = tbl2.rows[0].cells[i]; c.width = Cm(w)
    set_cell(c, h, bold=True, color=HDR_TXT, fill=HDR_FILL,
             align="left" if i==0 else "center")

# B=0 ABM baseline for highlighting
abm_Td_B0 = ABM_RESULTS["B=0  (no grandmothering)"]["Td"]
abm_PRLS_B0 = ABM_RESULTS["B=0  (no grandmothering)"]["PRLS"]

for ri, row in enumerate(rows_b):
    tr2 = tbl2.rows[ri+1]
    fill = ALT1 if ri%2==0 else ALT2
    abm = row
    aTr   = f'{abm["abm_Tr"]:.1f}±{abm["abm_TrSD"]:.2f}' if abm["abm_Tr"] else "—"
    aTd   = abm["abm_Td"]
    aPRLS = abm["abm_PRLS"]
    al70  = abm["abm_l70"]

    # Mark values below B=0 ABM baseline
    td_low   = (aTd is not None) and (aTd < abm_Td_B0) and (ri > 0)
    prls_low = (aPRLS is not None) and (aPRLS < abm_PRLS_B0) and (ri > 0)
    td_str   = f'{aTd:.1f} ↓' if td_low  else (f'{aTd:.1f}' if aTd else "—")
    prls_str = f'{aPRLS:.1f} ↓' if prls_low else (f'{aPRLS:.1f}' if aPRLS else "—")
    l70_str  = f'{al70:.2f}' if al70 else "—"

    vals = [
        row["condition"],
        f'{row["oc_gain"]:.2f}', str(row["oc_Tstar"]),
        f'{row["oc_PRLS"]:.1f}', f'{row["oc_l70"]:.2f}',
        aTr, td_str, prls_str, l70_str,
    ]
    for i, (v, w) in enumerate(zip(vals, CWID_B)):
        c = tr2.cells[i]; c.width = Cm(w)
        # Red for ↓ (below B=0 baseline); green for ↑ (above B=0 baseline)
        abm_col = None
        if i >= 5 and v != "—":
            if "↓" in str(v): abm_col = RED
            elif i == 6 and aTd and aTd > abm_Td_B0 and ri > 0: abm_col = GREEN
            elif i == 7 and aPRLS and aPRLS > abm_PRLS_B0 and ri > 0: abm_col = GREEN
        set_cell(c, v, color=abm_col, fill=fill,
                 align="left" if i==0 else "center")

doc.add_paragraph()
fn2 = doc.add_paragraph()
fn2.add_run(
    "Red ↓ : ABM value below the B=0 no-grandmothering ABM baseline — "
    "the Kirkwood indirect effect (grandmothering boosts daughter fecundity, "
    "accelerating daughter aging via δ(u), reducing population Td and PRLS "
    "before selection on grandmother longevity can respond).  "
    "Green: ABM value above B=0 baseline.  "
    "OC l₇₀ = 0.43 at B=1 MGM+PGM matches the Gurven–Kaplan target (0.40–0.47).  "
    "ABM l₇₀ at all conditions is far below OC because the selection coefficient "
    "s ≈ 0.001 is too small to drive maintenance evolution in tractable run times.").font.size = Pt(8)

doc.save("stepwise_table.docx")
print("Saved stepwise_table.docx")
